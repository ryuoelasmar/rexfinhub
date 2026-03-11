"""Generate a 1-page Word doc explaining REX email reports for executive review."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(0)

NAVY = RGBColor(0x1A, 0x1A, 0x2E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, hex_color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"}))


def _section_label(text, hex_color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _build_table(rows_data):
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    widths = [Inches(1.3), Inches(0.55), Inches(2.9), Inches(2.25)]

    for i, h in enumerate(["Report", "Cadence", "Description", "Key Metrics"]):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = WHITE
        _shade(c, "1A1A2E")

    for name, cadence, desc, metrics in rows_data:
        row = t.add_row()
        r = row.cells[0].paragraphs[0].add_run(name)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = NAVY

        r = row.cells[1].paragraphs[0].add_run(cadence)
        r.font.size = Pt(8)

        r = row.cells[2].paragraphs[0].add_run(desc)
        r.font.size = Pt(8)
        r.font.color.rgb = DARK

        r = row.cells[3].paragraphs[0].add_run(metrics)
        r.font.size = Pt(7.5)
        r.font.color.rgb = GRAY

    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)


# -- Title --
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("REX ETP Intelligence Platform")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Email Report Suite")
r.font.size = Pt(10)
r.font.color.rgb = GRAY
p.paragraph_format.space_after = Pt(4)

# -- Intro --
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run(
    "Automated email reports providing competitive intelligence across the leveraged, inverse, "
    "and income ETP landscape. Generated locally from SEC EDGAR and Bloomberg data, "
    "separate from the REX FinHub website. ETN data integration is in progress."
)
r.font.size = Pt(8.5)
r.font.color.rgb = DARK

# =========================================================================
_section_label("Finalized", "27AE60")

_build_table([
    (
        "Weekly ETP Report",
        "Weekly",
        "Full weekly wrap combining filing activity with Bloomberg data. "
        "REX performance by product suite, benchmarked against the broader universe.",
        "AUM & flows by suite; Winners/Losers/Yielders; "
        "Filing activity (7D); 5-category market landscape with issuer rankings"
    ),
    (
        "L&I Report",
        "Weekly",
        "Leveraged & inverse deep dive. Single-stock and index segments with "
        "3-year AUM trends, issuer share, and underlier breakdowns.",
        "Segment AUM & flows; REX share %; Issuer market share; "
        "AUM by underlier/category; Top 10 weekly inflows & outflows"
    ),
    (
        "Income Report",
        "Weekly",
        "Income/covered-call deep dive. Same structure as L&I with added "
        "yield analysis and traditional vs. synthetic strategy breakdown.",
        "Segment AUM & flows; REX share %; Avg yield by issuer; "
        "AUM by underlier with yield data; Top 10 weekly inflows & outflows"
    ),
    (
        "Flow Report",
        "Weekly",
        "Competitive flow positioning. Issuer-level, category-level, and "
        "fund-level analysis showing where capital is moving and REX's capture.",
        "Issuer flow rankings; Category flow analysis; "
        "Fund-level peer comparison; REX vs. competitor flows"
    ),
    (
        "Osprey Daily Report",
        "Daily",
        "Tracks Osprey fund performance: historical and current AUM, revenue, "
        "price, premium/discount, and key fund health metrics.",
        "AUM trend; Revenue; NAV & price; Premium/discount; "
        "Fund-specific performance metrics"
    ),
])

# =========================================================================
_section_label("Work in Progress", "E67E22")

_build_table([
    (
        "Daily ETP Report",
        "Daily",
        "Overnight filing surveillance. New SEC filings, Bloomberg-detected launches, "
        "upcoming effective dates, and REX market position snapshot.",
        "REX AUM & daily flows; Top movers; New filings by trust; "
        "Market landscape; Pending funds countdown"
    ),
    (
        "Monthly Asia Report",
        "Monthly",
        "Regional ETP intelligence covering leveraged and income product growth, "
        "issuer activity, and flow trends across Asia-Pacific.",
        "In development"
    ),
    (
        "Quarterly 13F Report",
        "Quarterly",
        "Institutional ownership from SEC 13F filings. Tracks which institutions "
        "are buying/selling leveraged and income ETPs and position changes.",
        "In development"
    ),
    (
        "Monthly Structured Product Report",
        "Monthly",
        "Covers all issued notes and the broader structured products landscape. "
        "Tracks issuance volume, product types, and market trends.",
        "In development"
    ),
])

# -- Footer --
doc.add_paragraph("")
p = doc.add_paragraph()
r = p.add_run("Sources: ")
r.bold = True
r.font.size = Pt(8)
r.font.color.rgb = DARK
r = p.add_run("SEC EDGAR (194 trusts) | Bloomberg (2,500+ ETPs)")
r.font.size = Pt(8)
r.font.color.rgb = GRAY

p = doc.add_paragraph()
r = p.add_run("Distribution: ")
r.bold = True
r.font.size = Pt(8)
r.font.color.rgb = DARK
r = p.add_run("ETFUpdates@rexfin.com + internal stakeholders, sent via REX FinHub admin panel.")
r.font.size = Pt(8)
r.font.color.rgb = GRAY

# -- Save --
out = Path(r"C:\Projects\rexfinhub\reports\REX_Email_Reports_Overview_v3.docx")
doc.save(str(out))
print(f"Saved to {out}")
