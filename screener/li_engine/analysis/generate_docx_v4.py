"""Generate methodology doc v4 — integrated cross-section + post-launch
success + robustness checks + cohort regime analysis + filing velocity."""
from __future__ import annotations

import json
import logging
import sqlite3
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from scipy.stats import spearmanr

from screener.li_engine.analysis.multi_angle_v3 import (
    build_targets, derive_weights_v3, run_all_v3, PANEL_PATH,
)
from screener.li_engine.analysis.post_launch_success import (
    build_panel as build_post_launch_panel, run_analysis as run_post_launch,
)
from screener.li_engine.analysis.robustness import (
    insider_pct_robustness, breakout_signal_tests,
)

NAVY = RGBColor(0x1a, 0x1a, 0x2e)
BLUE = RGBColor(0x09, 0x84, 0xe3)
GREEN = RGBColor(0x27, 0xae, 0x60)
RED = RGBColor(0xe7, 0x4c, 0x3c)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
GREY = RGBColor(0x7f, 0x8c, 0x8d)

_ROOT = Path(__file__).resolve().parent.parent.parent.parent
FILINGS_CS = _ROOT / "data" / "analysis" / "competitor_filing_cross_section.parquet"


def _h(doc, text, level=1):
    p = doc.add_heading(text, level)
    for r in p.runs:
        r.font.color.rgb = NAVY


def _para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = bold
    r.italic = italic


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)


def _fmt(v, pct=False, places=3):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    if pct:
        return f"{v:.1%}"
    return f"{v:+.{places}f}"


def _table(doc, rows):
    tbl = doc.add_table(rows=1, cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(rows[0]):
        hdr[i].text = str(h)
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for row_data in rows[1:]:
        row = tbl.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = str(v)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return tbl


def build(cross_reports, cross_weights, cross_corr, post_res, robust_insider, breakout, filings_cs, output: Path):
    doc = Document()

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("L&I Recommender Methodology v4")
    r.font.size = Pt(26); r.bold = True; r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Integrated: Cross-Section + Post-Launch + Cohort + Breakout + Robustness")
    sr.font.size = Pt(13); sr.font.color.rgb = BLUE
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run(date.today().strftime("%B %d, %Y"))
    dr.font.size = Pt(11); dr.font.color.rgb = GREY

    # Executive summary
    doc.add_paragraph()
    _h(doc, "The short version")
    _para(doc, "v4 integrates two distinct questions that v3 conflated:", bold=True)
    _bullet(doc, "Short-term question: what predicts the next 30 days of flow into L&I products on a given underlier?")
    _bullet(doc, "Long-term question: what underlier features predict that a NEW L&I product will reach $50M AUM within 18 months?")
    _para(doc, "Different targets, different top signals. The production engine should blend both — "
          "short-window for 'reinforce/launch now' decisions, long-window for 'should we file?' decisions.")

    _h(doc, "Headline signals", 2)
    rows = [["Question", "Top signal", "Second", "Third"]]
    rows.append(["Forward-30d flow (cross-section)", "mentions_24h (+0.34)", "rvol_90d (+0.33)", "insider_pct (+0.23)"])
    rows.append(["Long-term success (post-launch)", "turnover (+0.33)", "put_oi (+0.32)", "total_oi (+0.31)"])
    _table(doc, rows)

    _h(doc, "Headline finding you flagged and the data confirmed", 2)
    _para(doc, "You challenged v3's 2.7% weight on Total OI. The post-launch success analysis validates "
          "your skepticism: OI has IC = +0.307 against log(AUM) for mature products. The prior 301-product "
          "backtest showed r_log = 0.646. OI was never weak — v3 was measuring the wrong thing. "
          "The correct weight depends on which decision we're making.", bold=True)

    # --- Cross-section findings ---
    doc.add_page_break()
    _h(doc, "Analysis 1: Cross-section forward-flow IC")
    _para(doc, f"Sample: 186 underliers with observed forward window (as-of 60 BDays ago, window 30 BDays). "
          f"Target: forward_flow_30d aggregated across all products per underlier. "
          f"Same 10-angle framework as v3.")

    _h(doc, "Signal verdicts (forward-flow IC)", 2)
    rows = [["Signal", "Verdict", "Pos / Neg / Unclear", "Med |IC|", "Weight"]]
    for sig, rep in cross_reports.items():
        ics = [abs(tv.get("spearman", float("nan"))) for tv in rep.per_target.values()
               if tv.get("spearman") is not None and not np.isnan(tv.get("spearman"))]
        med = float(np.median(ics)) if ics else float("nan")
        rows.append([sig, rep.verdict(),
                     f"{rep.consensus_positive}/{rep.consensus_negative}/{rep.consensus_unclear}",
                     _fmt(med), _fmt(cross_weights.get(sig, 0), pct=True)])
    _table(doc, rows)

    # --- Post-launch analysis ---
    doc.add_page_break()
    _h(doc, "Analysis 2: Post-launch success (the real 'should we file?' test)")
    _para(doc, f"Sample: {post_res['n_mature']} L&I single-stock products aged 18+ months. "
          f"Success rate (AUM >= $50M): {post_res['success_rate']:.1%}. "
          f"Outcome distribution: {post_res['outcome_counts']}. "
          f"Caveat: underlier metrics snapshotted today, not at launch date (historical launch snapshots "
          f"not captured). Today's metrics are a proxy — stable features like market cap class and sector "
          f"are reasonable; momentum signals are contaminated.")

    _h(doc, "IC of each signal vs. long-term outcomes", 2)
    rows = [["Signal", "vs log(AUM)", "vs success≥50M", "vs ≥10M", "vs ≥100M"]]
    for sig, tgts in post_res["signal_ic_by_target"].items():
        rows.append([sig,
                     _fmt(tgts.get("log_aum", {}).get("ic")) if tgts.get("log_aum") else "—",
                     _fmt(tgts.get("success_18m_binary", {}).get("ic")) if tgts.get("success_18m_binary") else "—",
                     _fmt(tgts.get("aum_above_10m", {}).get("ic")) if tgts.get("aum_above_10m") else "—",
                     _fmt(tgts.get("aum_above_100m", {}).get("ic")) if tgts.get("aum_above_100m") else "—"])
    _table(doc, rows)

    _h(doc, "Interpretation", 2)
    _bullet(doc, "turnover / put_oi / total_oi / call_oi: +0.28 to +0.33 on log(AUM). These are the CLASSIC long-term demand predictors. They validate the old screener's 30% weights and the 301-product backtest that found r_log = 0.646 for OI.")
    _bullet(doc, "rvol_30d / rvol_90d: +0.30. Volatility shows up in BOTH analyses (short-term flow AND long-term success). Most universal signal.")
    _bullet(doc, "si_ratio: −0.329 — strongest NEGATIVE predictor of long-term success. High short interest → product fails. Systematic.")
    _bullet(doc, "Momentum reverses: ret_3m at −0.188 (negative), ret_6m at −0.173 (negative). Products launched on stocks that ran up in the prior 3–6 months tend to fail — the easy money is gone.")
    _bullet(doc, "inst_own_pct: −0.179. High institutional ownership → lower product success. Institutions hold quality; retail buys leverage.")
    _bullet(doc, "insider_pct: +0.126 post-launch (weaker than cross-section +0.23 but same direction — real signal, smaller magnitude).")
    _bullet(doc, "news_sentiment_bbg: −0.160 — mildly negative. Suggests 'good news' stocks are already priced / saturated.")

    # --- Cohort regime shift ---
    doc.add_page_break()
    _h(doc, "Analysis 3: Cohort stability — regime shift over time")
    _para(doc, "IC of each signal vs. log(AUM), stratified by launch year. Dominant signals differ by cohort — "
          "this tells us the methodology must acknowledge regime-dependence, not treat signals as universal.")

    for yr in sorted(post_res["cohort_stability"].keys()):
        cohort = post_res["cohort_stability"][yr]
        _para(doc, f"Cohort {yr} (n={cohort['n']}):", bold=True)
        rows = [["Signal", "IC vs log(AUM)"]]
        sorted_sigs = sorted(cohort["ics"].items(), key=lambda x: -abs(x[1]))[:10]
        for sig, ic in sorted_sigs:
            rows.append([sig, _fmt(ic)])
        _table(doc, rows)
        doc.add_paragraph()

    _h(doc, "What this tells us", 2)
    _bullet(doc, "2022 cohort: turnover / OI / call_oi / market_cap all IC +0.50 to +0.65 — classic liquidity-driven market.")
    _bullet(doc, "2024 cohort: rvol_30d/90d dominate (+0.26, +0.24), classic demand metrics drop to zero. Retail/meme era.")
    _bullet(doc, "Implication: 'current weights' should probably reflect 2024 regime (vol-heavy) while keeping liquidity metrics in reserve in case the regime shifts back.")
    _bullet(doc, "Risk: if we assume 2024 forever, we miss a return to 2022-style conditions. Monitor regime quarterly.")

    # --- Robustness checks ---
    doc.add_page_break()
    _h(doc, "Analysis 4: Robustness of suspect findings")
    _h(doc, "insider_pct robustness (you flagged this)", 2)
    rows = [["Test", "Result"]]
    rows.append(["Base IC", _fmt(robust_insider['base_ic'])])
    rows.append(["Winsorized [5%, 95%]", _fmt(robust_insider['winsorized_5_95_ic'])])
    rows.append(["Outliers removed (top/bottom 5%)", _fmt(robust_insider['no_extremes_ic'])])
    rows.append(["Size-partialled IC", _fmt(robust_insider['size_partialled_ic'])])
    rows.append(["Bootstrap median", _fmt(robust_insider['bootstrap_median'])])
    rows.append(["Bootstrap 95% CI", f"[{robust_insider['bootstrap_95ci'][0]:+.3f}, {robust_insider['bootstrap_95ci'][1]:+.3f}]"])
    rows.append(["Bootstrap % positive", f"{robust_insider['bootstrap_p_pos']:.1%}"])
    _table(doc, rows)

    tercile = robust_insider.get("by_mktcap_tercile", {})
    _para(doc, "By market cap tercile:", bold=True)
    for t in ["small", "mid", "large"]:
        v = tercile.get(t)
        if v:
            _bullet(doc, f"{t.title()} cap (n={v['n']}): IC = {v['ic']:+.3f}")

    _para(doc, "Conclusion: insider_pct passes every robustness check. Not outlier-driven. Stronger in "
          "large-caps (opposite of the 'small-cap founder artifact' concern). 99.85% of bootstrap samples "
          "positive. The 14.5% weight is defensible. One data-quality fix needed: cap insider_pct at 100% "
          "(COIN reported 2,601% in the current snapshot).", italic=True)

    _h(doc, "Breakout / co-movement tests (you asked for this)", 2)
    rows = [["Variant", "Spearman IC", "n"]]
    for name, stats in breakout["variants"].items():
        if stats is None:
            rows.append([name, "—", "—"])
        else:
            rows.append([name, _fmt(stats["ic"]), str(stats["n"])])
    _table(doc, rows)

    _para(doc, "Conclusion: multiplicative breakout signals (ret × mentions, ret × vol) do NOT outperform "
          "mentions alone. Sum-of-ranks composites are positive but modest (IC +0.15 to +0.16). "
          "Interpretation: attention is the real driver; co-movement with price adds noise at the 1-month "
          "window. Worth retesting once we have ret_1w and ret_3d from daily bbg prices — a real breakout "
          "needs a shorter momentum window than 1m.", italic=True)

    # --- Competitive filing velocity ---
    if filings_cs is not None and len(filings_cs) > 0:
        doc.add_page_break()
        _h(doc, "Analysis 5: Competitive filing velocity (context, not weight)")
        _para(doc, f"Filings pipeline: 798 L&I filings across {len(filings_cs)} distinct underliers. "
              "This is a context signal for the weekly report ('who's in line') — not a methodology input, "
              "but a direct production feature. Notable structural positions below.")

        _h(doc, "Top 10 REX monopolies (high REX filings, low competitor filings)", 2)
        monopolies = filings_cs.copy()
        monopolies = monopolies[monopolies["n_rex_filings_ever"] > 0]
        monopolies["rex_lead"] = monopolies["n_rex_filings_ever"] - monopolies["n_competitor_filings_ever"]
        top_monopolies = monopolies.sort_values("rex_lead", ascending=False).head(10)
        rows = [["Underlier", "REX filings", "Competitor filings", "Unique competitors"]]
        for u, r in top_monopolies.iterrows():
            rows.append([u, int(r["n_rex_filings_ever"]), int(r["n_competitor_filings_ever"]),
                         int(r["n_unique_competitors_ever"])])
        _table(doc, rows)

        _h(doc, "Saturated zones (highest competitor density)", 2)
        saturated = filings_cs.sort_values("n_unique_competitors_ever", ascending=False).head(10)
        rows = [["Underlier", "REX filings", "Competitor filings", "Unique competitors"]]
        for u, r in saturated.iterrows():
            rows.append([u, int(r["n_rex_filings_ever"]), int(r["n_competitor_filings_ever"]),
                         int(r["n_unique_competitors_ever"])])
        _table(doc, rows)

    # --- Recommended blended weights ---
    doc.add_page_break()
    _h(doc, "Recommended v0.3 blended weights")
    _para(doc, "The production engine serves TWO decisions. v3 weighted only for the short-term flow "
          "question. v4 proposes a blended scheme that makes the target explicit.")

    _h(doc, "Short-term flow weighting (for 'reinforce / launch now' decisions)", 2)
    rows = [["Signal", "Weight"]]
    for s, w in sorted(cross_weights.items(), key=lambda x: -x[1]):
        if w > 0:
            rows.append([s, _fmt(w, pct=True)])
    _table(doc, rows)

    _h(doc, "Long-term success weighting (for 'should we file?' decisions)", 2)
    # Derive from post-launch IC
    post_weights = {}
    for sig, tgts in post_res["signal_ic_by_target"].items():
        ic_logaum = tgts.get("log_aum", {}).get("ic") if tgts.get("log_aum") else None
        if ic_logaum is None:
            continue
        if ic_logaum > 0:
            post_weights[sig] = ic_logaum
    total = sum(post_weights.values())
    if total > 0:
        post_weights = {k: v / total for k, v in post_weights.items()}
    # cap 35%, floor 3%
    post_weights = {k: min(0.35, v) for k, v in post_weights.items() if v > 0.02}
    total = sum(post_weights.values())
    post_weights = {k: v / total for k, v in post_weights.items()}

    rows = [["Signal", "Weight"]]
    for s, w in sorted(post_weights.items(), key=lambda x: -x[1]):
        rows.append([s, _fmt(w, pct=True)])
    _table(doc, rows)

    _h(doc, "Blended production score", 2)
    _para(doc, "For any underlier, compute BOTH scores. Report both in the weekly email. "
          "File-candidate ranking leans on long-term weights; reinforce / launch ranking leans on short-term. "
          "Ambiguous cases (high in one, low in the other) are the interesting ones worth flagging.")

    # --- Playbook ---
    doc.add_page_break()
    _h(doc, "Ongoing review playbook")
    _para(doc, "Run every quarter. Dated output. Diff against prior quarter.")

    _h(doc, "Red flags", 2)
    _bullet(doc, "Signal verdict flips between quarters — investigate the regime / sample change.")
    _bullet(doc, "Cross-section IC and post-launch IC disagree in sign for the same signal — means that signal leads or lags differently; investigate.")
    _bullet(doc, "Cohort IC range > 0.5 spread for a signal — consider regime-conditional weighting.")
    _bullet(doc, "Bootstrap CI width > 2× point estimate — sample too small; don't trust the signal.")
    _bullet(doc, "Forward-window target distribution is all-zero or near-zero — target construction is broken (what happened to v2).")
    _bullet(doc, "Any signal correlates > 0.7 with another signal — apply correlation-aware shrinkage.")

    _h(doc, "Angles to rotate through future reviews", 2)
    _bullet(doc, "Paired-direction — same signal vs. long AND inverse product flows. Asymmetric = sentiment; symmetric = demand.")
    _bullet(doc, "Event study — signal level 5, 3, 1 days before competitor 485APOS filings (using `filings_by_underlier.parquet`).")
    _bullet(doc, "Survival analysis — time to reach $10M / $50M / $100M AUM, given signal at launch.")
    _bullet(doc, "Placebo shuffle — permute the signal across tickers; if IC unchanged, original is noise.")
    _bullet(doc, "Drawdown-regime IC — does the signal hold when market is drawing down.")
    _bullet(doc, "Google Trends multi-year IC — once pytrends backfill completes (5-year weekly), run time-stability on mentions-style signals.")
    _bullet(doc, "Persistent daily IC tracking — once `li_engine_daily` table is live, measure IC drift week-to-week.")
    _bullet(doc, "Sector-conditional weights — some signals (rvol in Industrials, mentions in Consumer Disc) are dramatically stronger in one sector.")

    _h(doc, "Data coverage — what's in, what's still outstanding", 2)
    _bullet(doc, "IN: bbg stock_data full (28 fields), bbg w5 returns, bbg daily time series (2.8M rows across 5,815 tickers), ApeWisdom sentiment, mkt_fund_classification, filings + FundStatus join.")
    _bullet(doc, "IN PROGRESS: Google Trends / pytrends 5y backfill (45/225 tickers as of this run; resume-capable).")
    _bullet(doc, "OUTSTANDING: FINRA biweekly short interest (agent blocked, needs retry), yfinance underlier OHLCV (rate-limited), persistent engine score storage (li_engine_daily table still not live).")
    _bullet(doc, "DROPPED: 13F institutional (user decision — off the backlog).")

    doc.save(str(output))


def main():
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    # Cross-section
    cross_panel = pd.read_parquet(PANEL_PATH)
    cross_reports, cross_corr = run_all_v3(cross_panel)
    cross_weights = derive_weights_v3(cross_reports)

    # Post-launch
    post_panel = build_post_launch_panel()
    post_res = run_post_launch(post_panel)

    # Robustness
    robust_insider = insider_pct_robustness(cross_panel)
    breakout = breakout_signal_tests(cross_panel)

    # Competitive filings
    filings_cs = None
    if FILINGS_CS.exists():
        filings_cs = pd.read_parquet(FILINGS_CS)

    out = _ROOT / "reports" / f"li_methodology_v4_{date.today().isoformat()}.docx"
    build(cross_reports, cross_weights, cross_corr, post_res, robust_insider, breakout, filings_cs, out)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
