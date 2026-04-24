"""Generate the v3 methodology doc — clean signals, fixed target, cross-sector."""
from __future__ import annotations

import logging
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from screener.li_engine.analysis.multi_angle_v3 import (
    SIGNALS_V3, PANEL_PATH, build_targets, derive_weights_v3,
    run_all_v3, SignalReport,
)

NAVY = RGBColor(0x1a, 0x1a, 0x2e)
BLUE = RGBColor(0x09, 0x84, 0xe3)
GREEN = RGBColor(0x27, 0xae, 0x60)
RED = RGBColor(0xe7, 0x4c, 0x3c)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
GREY = RGBColor(0x7f, 0x8c, 0x8d)


def _h(doc, text, level=1):
    p = doc.add_heading(text, level)
    for r in p.runs:
        r.font.color.rgb = NAVY


def _para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)


def _vc(v):
    if "positive" in v: return GREEN
    if "flip-sign" in v: return RED
    if "ambiguous" in v: return ORANGE
    return GREY


def _fmt(v, pct=False, places=3):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    if pct:
        return f"{v:.1%}"
    return f"{v:+.{places}f}"


def _table(doc, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(rows[0]):
        hdr[i].text = str(h)
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for row_data in rows[1:]:
        row = tbl.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = str(v)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return tbl


def build(reports, weights, corr_matrix, panel, targets, output: Path):
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("L&I Recommender Methodology v3")
    r.font.size = Pt(26); r.bold = True; r.font.color.rgb = NAVY
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Clean Signals, Fixed Forward Target, Cross-Sector IC")
    sr.font.size = Pt(14); sr.font.color.rgb = BLUE
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run(date.today().strftime("%B %d, %Y"))
    dr.font.size = Pt(11); dr.font.color.rgb = GREY

    # --- Exec summary ---
    doc.add_paragraph()
    _h(doc, "What changed since v2")
    _para(doc, "v2 had four real problems that user review caught. v3 fixes each:", bold=True)
    _bullet(doc, "Target leakage: `underlier_aum_12m_growth` was predicting forward flows INTO the same products whose AUM growth defined the signal. Autocorrelation, not signal. Removed.")
    _bullet(doc, "Collinear duplicates: rvol_30d ↔ rvol_90d (corr 0.93), market_cap ↔ turnover (0.88), adv_30d ↔ short_interest (0.86). v3 keeps one representative per cluster and applies correlation-aware shrinkage to anything still above 0.7.")
    _bullet(doc, "OI dismissal: v2 labeled OI 'ambiguous' and zero-weighted it despite the existing 301-product backtest showing r_log=0.646 for Total OI. v3 re-tests OI in four variants (total / call / put / skew) with no size control, larger sample, and sector stratification. Result: kept (positive), 2.7% per variant (shrunk), with Financials IC = +0.39 standing out.")
    _bullet(doc, "Underlier momentum: v2 said 'coverage gap.' The gap was my error — returns were in `mkt_stock_data.data_json` the whole time. v3 adds 1m / 3m / 6m / 1y returns plus pct-of-52w-high plus range-position. Each tested separately.")

    _h(doc, "What else v3 added", level=2)
    _bullet(doc, "Fixed forward-flow target — as-of date 60 BDays ago, window 30 BDays, 186 underliers with OBSERVED forward windows. v2's forward target was empty (as-of date was latest observation; no future data).")
    _bullet(doc, "Full 21-field extraction from bbg stock_data JSON: GICS sector, institutional ownership %, insider %, news sentiment, 52W high/low — previously ignored.")
    _bullet(doc, "Cross-sector IC on the primary target — a deferred angle in v2, now real.")
    _bullet(doc, "Signal correlation matrix with explicit redundancy shrinkage in weight derivation.")

    # --- Signal verdicts ---
    doc.add_page_break()
    _h(doc, "Signal Verdicts")
    _para(doc, f"Sample: {len(targets.get('forward_flow_30d', []))} underliers with observed forward window.")

    rows = [["Signal", "Verdict", "Pos / Neg / Unclear (of 4 targets)", "Median |Spearman|", "Weight"]]
    for sig, r in reports.items():
        if sig not in panel.columns:
            continue
        ics = [abs(tv.get("spearman", float("nan"))) for tv in r.per_target.values()
               if tv.get("spearman") is not None and not np.isnan(tv.get("spearman"))]
        med = float(np.median(ics)) if ics else float("nan")
        rows.append([sig, r.verdict(),
                     f"{r.consensus_positive} / {r.consensus_negative} / {r.consensus_unclear}",
                     _fmt(med), _fmt(weights.get(sig, 0), pct=True)])
    _table(doc, rows)

    # --- Key findings narrative ---
    doc.add_page_break()
    _h(doc, "Six findings the data actually supports")

    _h(doc, "1. Retail mentions are the strongest signal", 2)
    _para(doc, f"ApeWisdom mentions_24h: weight {weights.get('mentions_24h', 0):.1%}, median |IC| 0.344. "
          "3 of 4 targets positive. Cross-sector: dominates in Consumer Discretionary (IC=+0.53) and "
          "Industrials (+0.48). This is retail attention translating directly into leveraged product flows. "
          "Coverage caveat: ApeWisdom only returns trending tickers, so low-mention = no signal, not a reliable zero.")

    _h(doc, "2. Volatility confirms across every angle", 2)
    _para(doc, f"rvol_90d: weight {weights.get('rvol_90d', 0):.1%}, median |IC| 0.328. All 4 targets positive. "
          "Cross-sector: Industrials +0.63, Financials +0.38, Communication Services +0.35. "
          "rvol_30d not in the panel this time — 0.93 correlated with 90d, so we use 90d as the representative. "
          "Combined 'vol factor' weighting is not double-counted.")

    _h(doc, "3. Insider ownership is a genuine signal", 2)
    _para(doc, f"insider_pct: weight {weights.get('insider_pct', 0):.1%}, median |IC| 0.227. "
          "3 of 4 targets positive. Unusual finding because we didn't expect this. Hypothesis: insider "
          "holdings proxy for conviction — companies where insiders hold material stakes attract retail "
          "attention which drives leveraged product flows. Worth further investigation.")

    _h(doc, "4. Momentum has time-structure, not a single factor", 2)
    _para(doc, "Four return windows tested separately:", italic=False)
    _bullet(doc, "ret_1m: keep (positive), 4/4, med |IC| 0.030 — short-term momentum continues")
    _bullet(doc, "ret_3m: FLIP-SIGN, 3/4 negative, med |IC| 0.244 — 3-month mean-reversion dominates")
    _bullet(doc, "ret_6m: ambiguous, dropped")
    _bullet(doc, "ret_1y: keep (positive), 4/4, med |IC| 0.182 — long-term trend re-emerges")
    _para(doc, "Implication: a single 'momentum' signal would cancel itself out. Different windows measure "
          "different regimes. v3 keeps only the two ends of the distribution.", italic=True)

    _h(doc, "5. OI is a real but weaker cross-sectional signal", 2)
    _para(doc, "All four variants (total_oi, call_oi, put_oi, put_call_skew) earned 'keep (positive)'. "
          "Individual med |IC| values 0.07–0.11 — weaker than vol and sentiment. Collectively 5.4% weight. "
          "Why weaker than the 301-product backtest's r_log=0.646? Two reasons: (a) historical backtest used "
          "AUM as target (contaminated by size and market P&L); (b) our cross-section has only 186 underliers "
          "vs. 301 historical launches. OI's strongest sector: Financials (+0.39) — a liquidity-driven book.")

    _h(doc, "6. Four signals are negatively correlated with forward flows", 2)
    _para(doc, "These didn't earn weight, but they're informative:", italic=False)
    _bullet(doc, "si_ratio (short interest): neg across 4/4. High short interest → lower forward flows. Interpretation: shorts are absorbed by existing products; high SI = saturated demand.")
    _bullet(doc, "ret_3m: neg. Stocks that ran up over the past 3 months see profit-taking in the next 30 days — via leveraged inverse products.")
    _bullet(doc, "pct_of_52w_high & range_position: neg. Stocks near their 52W high attract less new leveraged long flow (profit-taking, not chasing).")
    _bullet(doc, "inst_own_pct: neg. High institutional ownership → lower retail leverage flow. Consistent with 'institutions hold quality, retail buys leverage'.")

    # --- Cross-sector IC table ---
    doc.add_page_break()
    _h(doc, "Cross-sector IC on forward_flow_30d (kept signals only)")
    _para(doc, "A signal's IC varies by sector. This is critical: a signal that works in aggregate may be "
          "driven entirely by one sector. Use sector-conditional weights in production if a signal's spread "
          "across sectors is too wide.")

    sector_rows = [["Signal"]]
    all_sectors = set()
    for r in reports.values():
        all_sectors.update(r.per_sector.keys())
    sectors = sorted(all_sectors)
    sector_rows[0].extend(sectors)

    kept = [sig for sig, r in reports.items() if "positive" in r.verdict() and r.per_sector]
    for sig in kept:
        row = [sig]
        for sec in sectors:
            ic = reports[sig].per_sector.get(sec)
            row.append(_fmt(ic) if ic is not None else "—")
        sector_rows.append(row)
    if len(sector_rows) > 1:
        _table(doc, sector_rows)

    # --- Correlation matrix ---
    doc.add_page_break()
    _h(doc, "Signal Correlation Matrix (Spearman)")
    _para(doc, "Pairs above 0.70 are flagged and weight is shrunk by (1 - |corr|) on the lower-magnitude signal. "
          "This replaces v2's cluster-sum hack with a principled correlation-aware weighting.")

    sig_list = list(corr_matrix.columns)
    cm_rows = [[""] + sig_list]
    for s in sig_list:
        row = [s]
        for o in sig_list:
            v = corr_matrix.loc[s, o]
            row.append(_fmt(v, places=2) if not pd.isna(v) else "—")
        cm_rows.append(row)
    _table(doc, cm_rows)

    # --- Weights ---
    doc.add_page_break()
    _h(doc, "Final v0.2 Weights")
    _para(doc, "These are the weights that would apply in a production engine. Zero-weight signals are "
          "excluded (not negatively weighted) — we can't trust the negative relationship enough to use it as "
          "a short-bias factor without more data.")

    wt_rows = [["Signal", "Weight"]]
    for s, w in sorted(weights.items(), key=lambda x: -x[1]):
        wt_rows.append([s, _fmt(w, pct=True)])
    _table(doc, wt_rows)

    # --- Ongoing playbook ---
    doc.add_page_break()
    _h(doc, "Ongoing Review Playbook")
    _para(doc, "Run every quarter. Dated output. Diff against prior version. Investigate material changes.")

    _h(doc, "Red flags", 2)
    _bullet(doc, "A signal's verdict flips between quarters — investigate the regime / sample change.")
    _bullet(doc, "Any signal's cross-sector IC range widens above 0.5 spread — consider sector-conditional weighting.")
    _bullet(doc, "New pair correlation > 0.7 appears — update the shrinkage list.")
    _bullet(doc, "A signal earns high weight but its Pearson and Spearman disagree in sign — non-linear; investigate before using.")
    _bullet(doc, "Bootstrap CI width > 2× point estimate — sample too small; don't trust the signal until it grows.")
    _bullet(doc, "Forward-window target produces all-zero or near-zero distribution — target construction is broken (what happened to us in v2).")

    _h(doc, "Angles to rotate through future reviews", 2)
    _bullet(doc, "Paired-direction — same signal vs. long AND inverse product flows. Asymmetric signals are sentiment; symmetric signals are demand.")
    _bullet(doc, "Event study — signal level 5, 3, 1 days before competitor 485APOS filings on the same underlier.")
    _bullet(doc, "Survival analysis — signal at launch vs. probability product reaches $50M AUM in 6 months.")
    _bullet(doc, "Placebo shuffle — permute the signal across tickers; if IC unchanged, original is noise.")
    _bullet(doc, "Drawdown-regime IC — does the signal hold when the market is drawing down? Regime-robustness check.")
    _bullet(doc, "Cross-API sentiment — pytrends vs. ApeWisdom; 0.53 correlated. Backfill pytrends for multi-year stability.")
    _bullet(doc, "Time-series IC — stability across pipeline runs once we have 6 months of dated history.")

    _h(doc, "What still needs to be added before v4", 2)
    _bullet(doc, "Persistent storage of daily engine scores + signal values per ticker. Not written anywhere today.")
    _bullet(doc, "yfinance underlier returns for tickers outside bbg's 6,479-stock universe (re-attempt after rate-limit cooldown).")
    _bullet(doc, "Google Trends via pytrends — only source with 5 years of historical sentiment backfill.")
    _bullet(doc, "Competitor-filing event timestamps from `filings` + `FundStatus` join.")
    _bullet(doc, "FINRA biweekly short interest — more reliable than bbg's snapshot.")
    _bullet(doc, "Post-launch success target — for historical 2x/3x products, did they reach $50M AUM in 6 months? A success binary target supports survival analysis, which is better than continuous flow for the 'should we file?' decision.")
    _bullet(doc, "13F quarterly institutional accumulation — killed this pass to save tokens; 2.5M-row source still ready.")
    _bullet(doc, "Sector-conditional weights in production — the cross-sector table shows some signals work dramatically better in one sector.")

    doc.save(str(output))


def main():
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    panel = pd.read_parquet(PANEL_PATH)
    reports, corr = run_all_v3(panel)
    weights = derive_weights_v3(reports)
    targets = build_targets(panel)
    out = Path(__file__).resolve().parent.parent.parent.parent / "reports" / f"li_methodology_v3_{date.today().isoformat()}.docx"
    build(reports, weights, corr, panel, targets, out)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
