"""Generate the v2 methodology Word doc from the expanded analysis."""
from __future__ import annotations

import logging
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from screener.li_engine.analysis.multi_angle_v2 import (
    SIGNALS_V2, PANEL_PATH, build_targets_v2, derive_weights_v2, run_all_v2,
    SignalReport,
)

NAVY = RGBColor(0x1a, 0x1a, 0x2e)
BLUE = RGBColor(0x09, 0x84, 0xe3)
GREEN = RGBColor(0x27, 0xae, 0x60)
RED = RGBColor(0xe7, 0x4c, 0x3c)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
GREY = RGBColor(0x7f, 0x8c, 0x8d)


def _h1(doc, text):
    p = doc.add_heading(text, 1)
    for r in p.runs:
        r.font.color.rgb = NAVY
def _h2(doc, text):
    p = doc.add_heading(text, 2)
    for r in p.runs:
        r.font.color.rgb = NAVY
def _h3(doc, text):
    p = doc.add_heading(text, 3)
    for r in p.runs:
        r.font.color.rgb = NAVY


def _para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)


def _verdict_color(v):
    if "positive" in v:
        return GREEN
    if "flip-sign" in v:
        return RED
    if "ambiguous" in v:
        return ORANGE
    return GREY


def _fmt(v, places=3, pct=False):
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    if pct:
        return f"{v:.1%}"
    return f"{v:+.{places}f}"


def _table(doc, data, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(data[0]))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(data[0]):
        hdr[i].text = str(h)
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
    for row_data in data[1:]:
        row = tbl.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = str(v)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return tbl


def _coverage_lines(panel: pd.DataFrame, groups: dict[str, list[str]]) -> list[list]:
    rows = [["Source", "Signals", "Coverage (ticker count)"]]
    for src, sigs in groups.items():
        present = [s for s in sigs if s in panel.columns]
        if not present:
            continue
        cov = {s: panel[s].notna().sum() for s in present}
        rows.append([src, ", ".join(present), f"range {min(cov.values())}–{max(cov.values())}"])
    return rows


def build(reports: dict[str, SignalReport], weights: dict[str, float],
          panel: pd.DataFrame, targets: dict[str, pd.Series], output: Path):

    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("L&I Recommender Methodology v2")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Expanded Multi-Angle Analysis + Forward-Flow Target")
    sr.font.size = Pt(14)
    sr.font.color.rgb = BLUE
    ds = doc.add_paragraph()
    ds.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = ds.add_run(date.today().strftime("%B %d, %Y"))
    dr.font.size = Pt(11)
    dr.font.color.rgb = GREY

    # Exec summary
    doc.add_paragraph()
    _h1(doc, "Executive Summary")
    _para(doc, "This version expands the v1 analysis (8 signals, contemporaneous target) to 17 signals "
          "across 4 data sources and adds a genuine forward-flow target from Bloomberg's daily "
          "time-series sheets — the kind of predictive IC the v1 run could not produce.")

    keepers = [s for s, r in reports.items() if "positive" in r.verdict()]
    flippers = [s for s, r in reports.items() if "flip-sign" in r.verdict()]
    ambigs = [s for s, r in reports.items() if "ambiguous" in r.verdict()]
    insuf = [s for s, r in reports.items() if r.verdict() == "insufficient-data"]

    _para(doc, "Top findings:", bold=True)
    _bullet(doc, f"Robust signals ({len(keepers)}): {', '.join(keepers)}")
    _bullet(doc, f"Size-bias / sign-flip ({len(flippers)}): {', '.join(flippers)}")
    _bullet(doc, f"Ambiguous / zero weight ({len(ambigs)}): {', '.join(ambigs)}")
    _bullet(doc, f"Insufficient data — coverage gap, not a verdict ({len(insuf)}): {', '.join(insuf)}")

    # Top weights table
    _para(doc, "Recommended v0.2 weights (consensus-driven):", bold=True)
    w_rows = [["Signal", "Weight"]]
    for s, w in sorted(weights.items(), key=lambda x: -x[1]):
        if w > 0:
            w_rows.append([s, _fmt(w, pct=True)])
    _table(doc, w_rows)

    # Data sources + coverage
    doc.add_page_break()
    _h1(doc, "Data Sources & Coverage")
    _para(doc, "What was used, what was tested and rejected, what was intended but blocked.")

    _h2(doc, "Used in this analysis")
    groups = {
        "Bloomberg stock_data (current snapshot)": [
            "market_cap", "adv_30d", "turnover_30d", "total_oi", "put_call_skew",
            "realized_vol_30d", "realized_vol_90d", "short_interest",
        ],
        "Bloomberg w5 sheet (cross-sectional returns)": [
            "ret_5d", "ret_1m", "ret_3m", "ret_6m", "ret_ytd", "ret_1y",
        ],
        "Bloomberg aum_history_json (12-month AUM trend on underlier's products)": [
            "underlier_aum_12m_growth",
        ],
        "ApeWisdom API (current retail-mention snapshot)": [
            "mentions_24h", "mentions_delta_24h",
        ],
    }
    _table(doc, _coverage_lines(panel, groups))

    _para(doc, "", )
    _para(doc, "Targets built from Bloomberg daily time series:", bold=True)
    tg_rows = [["Target", "n", "Description"]]
    tg_desc = {
        "forward_30d_flow": "Forward 30-day cumulative net flow per underlier, from daily data_flow sheet. First proper predictive target.",
        "raw_3m_flow": "Trailing 3-month net flow (contemporaneous).",
        "rank_3m_flow": "Percentile rank of 3-month flow — size invariant.",
        "log_3m_flow": "Sign-preserving log transform of 3-month flow.",
        "aum_growth": "Flow / starting AUM (known size-biased; kept for contrast).",
        "flow_adj_perf": "3-month flow minus estimated market P&L (decomposes Ryu's 'AUM = flow × perf' insight).",
    }
    for tn, t in targets.items():
        tg_rows.append([tn, str(len(t)), tg_desc.get(tn, "")])
    _table(doc, tg_rows)

    _h2(doc, "Tested and rejected")
    _bullet(doc, "StockTwits trending — HTTP 403 Cloudflare; no unauthenticated access in 2026.")
    _bullet(doc, "QuiverQuant WSB — HTTP 401; free tier discontinued, now bearer-token / paid.")
    _bullet(doc, "Finnhub social-sentiment — no FINNHUB_API_KEY in env; not tested this pass. GitHub #557 reports 403s on free tier even with a key.")
    _bullet(doc, "Reddit PRAW — requires credentials; none set.")

    _h2(doc, "Intended but blocked this pass")
    _bullet(doc, "yfinance — rate-limited even on 3-ticker probes after the initial batch. Would add RSI, 52w-high %, SPY-relative capture, up-streak, drawdown. Retry outside market hours.")
    _bullet(doc, "Google Trends (pytrends) — works but not yet plumbed in; ApeWisdom correlation 0.53 means it adds complementary signal, not redundant. Priority next pass because it gives real historical backfill.")
    _bullet(doc, "13F institutional signals — killed by user to conserve tokens; 2.5M-row database remains untouched.")
    _bullet(doc, "SEC filing-velocity signals — 613K filings table, not yet wired into the panel.")
    _bullet(doc, "Cross-sector labels — mkt_fund_classification has sector/theme columns; can stratify IC by sector next pass.")

    _h2(doc, "Coverage caveat that matters")
    _para(doc, "The w5 sheet provides per-ticker returns, but it tracks PRODUCT tickers (TSLT, NVDU, AAPX), "
          "not UNDERLIER tickers (TSLA, NVDA, AAPL). Only 8–9 underliers in our analytical universe also appear in w5. "
          "This is why ret_1m, ret_3m, ret_ytd show 'insufficient-data' in the verdict table — not a failure of the signal, "
          "a coverage mismatch. Underlier-level momentum must come from yfinance or a separate Bloomberg pull keyed on "
          "underlier tickers.", italic=True)

    # Methodology / angles
    doc.add_page_break()
    _h1(doc, "The Ten Angles (unchanged from v1)")
    angles = [
        ("Spearman rank-IC", "Monotonic, outlier-robust."),
        ("Pearson correlation", "Linear strength. Contrast with Spearman reveals curvature."),
        ("Quintile spread", "Top-20% minus bottom-20% outcome. The money test."),
        ("Size-controlled regression", "Signal coefficient after partialling out log market cap. Catches size-bias."),
        ("Mutual information", "Any relationship, linear or not. Catches U-shapes and thresholds."),
        ("Random-forest feature importance", "Joint contribution including interaction effects."),
        ("Lasso survival", "L1 regularization auto-selects features under sparsity."),
        ("Bootstrap confidence interval", "How certain are we the IC isn't noise."),
        ("Time stability (deferred)", "IC across pipeline runs. Requires dated history we are now accumulating."),
        ("Cross-sector consistency (deferred)", "IC within each sector. Requires attaching sector labels."),
    ]
    for name, desc in angles:
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"{name}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    _para(doc, "A signal must score positive across ≥60% of signed angles, in ≥60% of the 6 targets, to earn 'keep'. "
          "Signals that are ambiguous OR flip-sign OR insufficient-data get zero weight. "
          "We don't dilute weight on coin flips.", italic=True)

    # Signal verdicts
    doc.add_page_break()
    _h1(doc, "Signal Verdicts")
    v_rows = [["Signal", "Verdict", "Pos / Neg / Unclear (of 6 targets)", "Median |Spearman|", "Weight"]]
    for sig, r in reports.items():
        ics = [abs(tv["spearman"]) for tv in r.per_angle.values()
               if tv.get("spearman") is not None and not np.isnan(tv["spearman"])]
        med_ic = np.median(ics) if ics else float("nan")
        v_rows.append([
            sig, r.verdict(),
            f"{r.consensus_positive} / {r.consensus_negative} / {r.consensus_unclear}",
            _fmt(med_ic),
            _fmt(weights.get(sig, 0.0), pct=True),
        ])
    _table(doc, v_rows)

    # Key signal narratives
    _h1(doc, "What the data says")

    _h3(doc, "1. Underlier AUM growth dominates")
    _para(doc, f"`underlier_aum_12m_growth` is the top-weighted signal at {weights.get('underlier_aum_12m_growth', 0):.1%}. "
          "It measures the 12-month change in total AUM across all leveraged products on a given underlier, as a percentage. "
          "The logic: if the existing product book on an underlier is growing, demand is demonstrably live. "
          "This is the most honest demand signal in the panel because it observes actual behaviour, not proxies.")

    _h3(doc, "2. Retail sentiment (ApeWisdom) validates as a real signal")
    _para(doc, f"`mentions_24h` scores {weights.get('mentions_24h', 0):.1%}. 4 of 5 targets positive. "
          "But `mentions_delta_24h` flipped sign — mention spikes don't predict forward flows. "
          "Takeaway: absolute retail attention is the signal; day-over-day spikes are noise.")

    _h3(doc, "3. Volatility confirms from v1")
    combined_vol = weights.get("realized_vol_30d", 0) + weights.get("realized_vol_90d", 0)
    _para(doc, f"Realized vol at 30d and 90d both earn 'keep' across 5 of 5 targets. Combined weight: {combined_vol:.1%}. "
          "With cluster-aware scaling, each is capped to prevent double-counting the same information.")

    _h3(doc, "4. Market cap is a size-bias artifact, not a signal")
    _para(doc, "`market_cap` was negative across all 5 applicable targets — identical to v1. "
          "The artifact: flow/AUM mechanically favours smaller products; raw flow is correlated with size. "
          "We treat market cap as zero weight rather than flipping it into a negative factor.")

    _h3(doc, "5. OI, ADV, short interest did not earn weight")
    _para(doc, "These were either ambiguous or flipped sign depending on the target. "
          "Possible reasons: collinear with market cap; thinner population in the forward-flow subset; "
          "or genuinely uninformative at this sample size. Zero weight for now; re-test when sample grows.")

    _h3(doc, "6. Underlier momentum signals are a coverage gap, not a null result")
    _para(doc, "All w5 ret_* signals showed 'insufficient-data' because w5 carries product returns, "
          "not underlier returns. Closing this gap requires yfinance (currently rate-limited) or a "
          "dedicated Bloomberg underlier-returns pull.")

    # Ongoing review playbook
    doc.add_page_break()
    _h1(doc, "Ongoing Review Playbook")
    _para(doc, "Run this analysis every quarter. Record the output with a date stamp so we can watch signal stability drift.")

    _h2(doc, "Red flags — investigate, don't accept")
    _bullet(doc, "A signal's verdict flips between quarters (keep → ambiguous or vice versa).")
    _bullet(doc, "Spearman and Pearson disagree in sign — relationship is non-linear; understand before using.")
    _bullet(doc, "Size-controlled β has the opposite sign from raw IC — size is a confounder.")
    _bullet(doc, "Bootstrap CI width > 2× point estimate — we don't actually know what the signal is.")
    _bullet(doc, "Random-forest importance is high but Spearman IC near zero — non-linear / interaction effect unnamed.")
    _bullet(doc, "Quintile spread opposite sign from Spearman — relationship is non-monotonic (middle matters, not extremes).")
    _bullet(doc, "Forward-IC and contemporaneous-IC disagree — signal either leads or lags; declare which.")

    _h2(doc, "Angles to rotate through future reviews")
    _bullet(doc, "Time stability — decompose IC into trend + cyclical once we have ≥6 months of dated history (we're accumulating now).")
    _bullet(doc, "Cross-sector IC — stratify by sector / theme from mkt_fund_classification. A signal that works in semis but not biotech isn't universal.")
    _bullet(doc, "Event study around competitor 485APOS filings — does our signal spike on underliers JUST before competitors file?")
    _bullet(doc, "Survival analysis — signals at launch vs. probability of reaching $50M AUM in 6 months.")
    _bullet(doc, "Paired-direction — same signal vs. long AND inverse product flows. Direction-asymmetric signals are sentiment; symmetric are demand.")
    _bullet(doc, "Placebo shuffle — randomise the signal across tickers and re-run. If IC is similar, the original is noise.")
    _bullet(doc, "Drawdown regime — IC in drawdown periods vs. normal. Signals that break under stress aren't investable.")
    _bullet(doc, "Cross-API sentiment — pytrends vs. ApeWisdom; correlation 0.53 means both carry unique signal. Once backfilled, run both through multi-angle.")

    _h2(doc, "What must be in place before the next review")
    _bullet(doc, "Persistent storage of daily engine scores + signal values per ticker — currently not written anywhere. Highest priority.")
    _bullet(doc, "Sector / theme labels per underlier (mkt_fund_classification join).")
    _bullet(doc, "Competitor-filing event timestamps (filings + FundStatus join).")
    _bullet(doc, "yfinance underlier OHLCV backfill — reattempt outside market hours.")
    _bullet(doc, "Google Trends backfill via pytrends — 5 years of historical relative-interest for every ticker.")
    _bullet(doc, "13F institutional accumulation signal — 2.5M-row source ready; one query away.")
    _bullet(doc, "FINRA biweekly short-interest for backfill — SEC-exposed short interest is more reliable than bbg's snapshot.")

    _h2(doc, "Minimum-viable quarterly run")
    _para(doc, "Every quarter, script the full run: (1) rebuild expanded_signal_panel.parquet, (2) run multi_angle_v2, "
          "(3) generate a dated docx, (4) diff the new weights against the prior quarter's file. "
          "Flag any signal whose weight moves > 5 percentage points, any verdict that changes, "
          "any new 'insufficient-data' result. Those are the items for deep dive.")

    doc.save(str(output))


def main():
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    panel = pd.read_parquet(PANEL_PATH)
    reports = run_all_v2(panel)
    weights = derive_weights_v2(reports)
    targets = build_targets_v2(panel)
    out = Path(__file__).resolve().parent.parent.parent.parent / "reports" / f"li_methodology_v2_{date.today().isoformat()}.docx"
    build(reports, weights, panel, targets, out)
    print(f"Generated: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
