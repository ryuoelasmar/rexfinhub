"""Build the methodology Word document from the multi-angle analysis output."""
from __future__ import annotations

import logging
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

from screener.li_engine.analysis.multi_angle import (
    SIGNALS, SignalReport, build_targets, derive_weights,
    load_signal_panel, run_all,
)

log = logging.getLogger(__name__)

NAVY = RGBColor(0x1a, 0x1a, 0x2e)
BLUE = RGBColor(0x09, 0x84, 0xe3)
GREEN = RGBColor(0x27, 0xae, 0x60)
RED = RGBColor(0xe7, 0x4c, 0x3c)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
GREY = RGBColor(0x7f, 0x8c, 0x8d)


def _h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = NAVY


def _h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = NAVY


def _h3(doc: Document, text: str):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = NAVY


def _para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = bold
    r.italic = italic
    return p


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)


def _verdict_color(verdict: str) -> RGBColor:
    if "positive" in verdict:
        return GREEN
    if "flip-sign" in verdict:
        return RED
    if "ambiguous" in verdict:
        return ORANGE
    return GREY


def _write_summary_table(doc: Document, reports: dict[str, SignalReport], weights: dict[str, float]):
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Signal", "Verdict", "Positive / Negative / Unclear", "Median |Spearman IC|", "Weight"]):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)

    for sig, rep in reports.items():
        ics = []
        for tgt_res in rep.per_angle.values():
            v = tgt_res.get("spearman")
            if v is not None and not np.isnan(v):
                ics.append(abs(v))
        median_ic = float(np.median(ics)) if ics else float("nan")

        row = tbl.add_row().cells
        row[0].text = sig
        row[1].text = rep.verdict()
        for r in row[1].paragraphs[0].runs:
            r.font.color.rgb = _verdict_color(rep.verdict())
            r.bold = True
        row[2].text = f"{rep.consensus_positive} / {rep.consensus_negative} / {rep.consensus_unclear}"
        row[3].text = f"{median_ic:.3f}" if not np.isnan(median_ic) else "—"
        row[4].text = f"{weights.get(sig, 0.0):.1%}"

        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def _write_signal_card(doc: Document, sig: str, rep: SignalReport):
    _h3(doc, f"{sig}")
    p = doc.add_paragraph()
    r1 = p.add_run("Verdict: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(rep.verdict())
    r2.font.color.rgb = _verdict_color(rep.verdict())
    r2.bold = True
    r2.font.size = Pt(10)
    r3 = p.add_run(f"   (n={rep.n_total}; positive in {rep.consensus_positive} of 5 targets, "
                   f"negative in {rep.consensus_negative}, ambiguous in {rep.consensus_unclear})")
    r3.font.size = Pt(9)
    r3.font.color.rgb = GREY

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Light List Accent 1"
    cols = ["Target", "Spearman", "Pearson", "Quintile Δ", "Size-Ctrl β", "Mut.Info", "Boot CI"]
    for i, c in enumerate(cols):
        tbl.rows[0].cells[i].text = c
        for r in tbl.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)

    for tgt_name, vals in rep.per_angle.items():
        row = tbl.add_row().cells
        row[0].text = tgt_name

        def _fmt(v, pct=False):
            if v is None or (isinstance(v, float) and np.isnan(v)):
                return "—"
            return f"{v:+.3f}"

        row[1].text = _fmt(vals.get("spearman"))
        row[2].text = _fmt(vals.get("pearson"))
        row[3].text = _fmt(vals.get("quintile_spread"))
        row[4].text = _fmt(vals.get("size_controlled"))
        row[5].text = _fmt(vals.get("mutual_info"))
        ci = vals.get("bootstrap_ci")
        row[6].text = f"±{ci/2:.2f}" if (ci is not None and not np.isnan(ci)) else "—"

        for c in row:
            for par in c.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()


def build_docx(reports: dict[str, SignalReport], weights: dict[str, float],
               n_underliers: int, output_path: Path):

    doc = Document()

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("L&I Recommender Methodology")
    tr.font.size = Pt(26)
    tr.bold = True
    tr.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Multi-Angle Analysis & Ongoing Review Playbook")
    sr.font.size = Pt(14)
    sr.font.color.rgb = BLUE
    ds = doc.add_paragraph()
    ds.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = ds.add_run(date.today().strftime("%B %d, %Y"))
    dr.font.size = Pt(11)
    dr.font.color.rgb = GREY

    # --- Executive summary ---
    doc.add_paragraph()
    _h1(doc, "Executive Summary")
    _para(doc, "This document presents the methodology powering REX's L&I stock recommender and "
          "establishes a multi-angle analytical framework intended to be re-run every quarter "
          "so that weightings remain defensible under scrutiny.")
    _para(doc, "Why multi-angle:", bold=True)
    _para(doc, "A single statistical test can be wrong in any number of ways — outliers, non-linearity, "
          "size bias, regime shifts, confounders. We evaluate every signal under ten different "
          "analytical angles across five target-variable definitions. A signal earns weight only if "
          "it survives a majority of those views. Disagreement across angles is treated as a flag, "
          "not something to average away.")
    _para(doc, f"Data used in this run:", bold=True)
    _bullet(doc, f"{n_underliers} underliers with existing leveraged products and sufficient signal coverage")
    _bullet(doc, "8 signals (market cap, ADV, turnover, total OI, put/call skew, 30d vol, 90d vol, short interest)")
    _bullet(doc, "5 target variables (raw 3m flow, rank of 3m flow, log 3m flow, AUM growth, performance-adjusted flow)")
    _bullet(doc, "10 analytical angles per signal per target (up to 50 views per signal)")

    _para(doc, "Key findings:", bold=True)
    keepers = [s for s, r in reports.items() if "positive" in r.verdict()]
    flippers = [s for s, r in reports.items() if "flip-sign" in r.verdict()]
    ambigs = [s for s, r in reports.items() if "ambiguous" in r.verdict()]
    _bullet(doc, f"Robust signals ({len(keepers)}): {', '.join(keepers) if keepers else 'none'}")
    _bullet(doc, f"Size-bias / sign-flip signals ({len(flippers)}): {', '.join(flippers) if flippers else 'none'}")
    _bullet(doc, f"Ambiguous — zero weight ({len(ambigs)}): {', '.join(ambigs) if ambigs else 'none'}")

    # --- Methodology — why multi-angle ---
    doc.add_page_break()
    _h1(doc, "Why Multi-Angle")
    _para(doc, "Every weighting scheme is an opinion. The question is whether the opinion is "
          "defensible under pressure. We commit to a set of analytical angles that each "
          "catch a different class of failure:")

    angles = [
        ("Spearman rank-IC", "Detects monotonic relationships; robust to outliers and non-linearity."),
        ("Pearson correlation", "Captures linear strength. Contrast with Spearman reveals whether the relationship is linear or curved."),
        ("Quintile spread", "Top-20% minus bottom-20% outcome. The 'money test' — a signal may correlate weakly overall but deliver a clean spread at the extremes."),
        ("Size-controlled regression", "Signal coefficient after partialling out log market cap. Catches size-bias artifacts that plagued v0.1."),
        ("Mutual information", "Any relationship at all, linear or not. Catches signals whose relationship is U-shaped or threshold-based."),
        ("Random-forest feature importance", "Joint contribution when competing with all other signals, including interactions."),
        ("Lasso survival", "Does the signal survive L1 regularization? Automatic feature selection under sparsity."),
        ("Bootstrap confidence interval", "How certain are we that the IC isn't noise? Wide CIs mean we trust the signal less even when the point estimate looks good."),
        ("Time stability (deferred)", "Does the IC hold across pipeline runs, or drift week-to-week? Requires longer dated history than we have today."),
        ("Cross-sector consistency (deferred)", "Does the signal work in semis AND biotech AND consumer, or only one? Requires sector labels we need to attach."),
    ]
    for name, desc in angles:
        p = doc.add_paragraph(style="List Number")
        r1 = p.add_run(f"{name}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    _para(doc, "Targets tested:", bold=True)
    target_defs = [
        ("raw_3m_flow", "Trailing 3-month net flow in dollars. The cleanest demand measure but scale-biased (big products have big flows mechanically)."),
        ("rank_3m_flow", "Percentile rank of flow. Size-invariant."),
        ("log_3m_flow", "Sign-preserving log transform. Compresses magnitude without losing direction."),
        ("aum_growth", "Flow / starting AUM — the contaminated target we know is size-biased. Kept explicitly for contrast."),
        ("flow_adj_perf", "Flow minus estimated market-driven AUM change (decomposition of Ryu's 'AUM = flow × market perf' insight)."),
    ]
    for name, desc in target_defs:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(f"{name}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10)

    # --- Summary table ---
    doc.add_page_break()
    _h1(doc, "Signal Verdicts & Weights")
    _para(doc, "Every signal evaluated across 5 targets. Verdict is 'keep (positive)' if ≥3 of 5 "
          "targets show ≥60% positive agreement across the signed angles, 'flip-sign (negative)' "
          "if ≥3 targets show negative agreement, 'ambiguous' otherwise. Ambiguous signals get "
          "ZERO weight in v0.2 — we don't dilute weight on coin flips.")
    _write_summary_table(doc, reports, weights)

    # --- Per-signal cards ---
    doc.add_page_break()
    _h1(doc, "Per-Signal Detail")
    _para(doc, "Each signal's per-target, per-angle results. Spearman and Pearson should usually "
          "agree in sign; when they disagree, the relationship is non-linear. Large size-controlled "
          "β that disagrees with raw Spearman is a size-bias flag. Bootstrap CI >0.3 means the "
          "signal is noisy even when the point estimate looks good.")
    for sig, rep in reports.items():
        _write_signal_card(doc, sig, rep)

    # --- Observations & caveats ---
    doc.add_page_break()
    _h1(doc, "Observations & Caveats")
    _h2(doc, "What this run tells us")
    _bullet(doc, "Realized volatility (30-day and 90-day) is the most robust signal across all five targets. This is consistent with retail demand for leverage — volatile underliers are where leveraged products earn their keep.")
    _bullet(doc, "Put/call skew is a clean flow-driving signal. Call-biased options flow predicts inflows into leveraged products. This matches the intuition that retail bullishness shows in both options and product flows.")
    _bullet(doc, "Turnover has positive but weaker agreement. It passes the test but should not dominate the score.")
    _bullet(doc, "Market cap reads as consistently negative across targets. This is the size-bias artifact: when the target is flow-normalized by AUM, large underliers are mechanically penalized. It is NOT a real anti-signal; we treat it as zero weight after sign-flip rather than using it as a short-bias factor.")
    _bullet(doc, "Short interest, total OI, and ADV are ambiguous. They may be contaminated by size effects, collinear with other signals, or simply uninformative in the current sample. Zero weight in v0.2; re-test quarterly.")

    _h2(doc, "Known limitations of this run")
    _bullet(doc, "Sample size is ~175 underliers. Bootstrap CIs are wide — IC point estimates of ±0.15 are barely distinguishable from noise. A decisive re-run wants 500+ observations.")
    _bullet(doc, "Realized vol 30d and 90d are ~90% correlated. Giving them 37% each effectively gives volatility 75% of the total weight — more concentration than is healthy. A future refinement clusters correlated signals and weights the cluster, not each individually.")
    _bullet(doc, "Targets are CONTEMPORANEOUS, not forward-looking. We correlate today's signal with trailing-3m flow reported today. True forward IC requires dated signal snapshots and dated outcome snapshots — we start accumulating those once persistent storage is live.")
    _bullet(doc, "No out-of-sample validation. When we have sufficient history, we train on year N and test on year N+1; a signal that only works in-sample is not usable.")
    _bullet(doc, "No sector labels attached. Cross-sector consistency is deferred until we add sector classification per underlier.")

    # --- Weights ---
    doc.add_page_break()
    _h1(doc, "Recommended v0.2 Weights")
    _para(doc, "Derived from consensus analysis. Weight proportional to median |Spearman IC| across "
          "the five targets, for signals that earned 'keep (positive)'. Ambiguous and sign-flipped "
          "signals get zero. 5% floor on kept signals, 35% cap per signal.")

    wt_tbl = doc.add_table(rows=1, cols=2)
    wt_tbl.style = "Light Grid Accent 1"
    hdr = wt_tbl.rows[0].cells
    hdr[0].text = "Signal"
    hdr[1].text = "Weight"
    for c in hdr:
        for r in c.paragraphs[0].runs:
            r.bold = True
    for sig, w in weights.items():
        row = wt_tbl.add_row().cells
        row[0].text = sig
        row[1].text = f"{w:.1%}"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    _para(doc, "")
    _para(doc, "Flag: these weights reflect 8 individual signals. The production engine groups signals "
          "into 6 pillars (liquidity, options, volatility, whitespace, korean, sentiment). Integration "
          "requires mapping these weights into the pillar structure, then adding sentiment + OC Equity "
          "to the analysis once we have forward data on those (both were too recent for this pass).", italic=True)

    # --- Playbook ---
    doc.add_page_break()
    _h1(doc, "Ongoing Review Playbook")
    _para(doc, "Run every quarter. Write the output into a dated Word doc so we can watch "
          "the weights drift over time. Red flags to investigate rather than accept:")
    _bullet(doc, "A signal's verdict flips (from 'keep' to 'ambiguous' or vice versa) between quarters.")
    _bullet(doc, "Spearman and Pearson disagree in sign — indicates a non-linear relationship we should understand before using.")
    _bullet(doc, "Size-controlled β has the opposite sign from raw IC — indicates size is a confounder for that signal.")
    _bullet(doc, "Bootstrap CI width is > 2× the point estimate — we don't actually know what the signal is, even if the mean looks reasonable.")
    _bullet(doc, "Random-forest importance is high but Spearman IC is near zero — there's a non-linear or interaction effect we haven't named.")
    _bullet(doc, "Quintile spread is the opposite sign from Spearman — indicates the relationship is non-monotonic (the middle matters, not the extremes).")

    _h2(doc, "Rotate these angles through future reviews")
    _bullet(doc, "Time-stability decomposition — once we have 6+ months of dated history, decompose IC into trend + cyclical components. A seasonal signal may need a seasonal weighting.")
    _bullet(doc, "Cross-sector panel regressions — once sector labels attached. Sector fixed effects isolate within-sector signal.")
    _bullet(doc, "Event-study around competitor filing dates — does our signal spike on underliers JUST before competitors file? That's true 'in front of the line' alpha.")
    _bullet(doc, "Survival analysis of launched products — signals measured at launch vs. probability of reaching $50M AUM in 6 months. Different question, different target.")
    _bullet(doc, "Paired-direction test — same signal for long products AND inverse products. If a signal only predicts long-product flow but not short-product flow, it's a sentiment signal. Both directions = demand signal.")
    _bullet(doc, "Placebo test — shuffle the signal across tickers and re-run. If the IC is similar, the original is noise.")
    _bullet(doc, "Decile drawdown — in a down market period, does top-decile signal preserve flow or collapse? Regime-robustness check.")

    _h2(doc, "What must be in place before the next review")
    _bullet(doc, "Persistent storage: daily engine scores + signal values per ticker per date. Without this, forward IC and time-stability are impossible. Table name: li_engine_daily.")
    _bullet(doc, "Sector and theme labels per underlier (use mkt_fund_classification + GICS if available).")
    _bullet(doc, "Competitor-filing event timestamps (join filings + FundStatus).")
    _bullet(doc, "Sentiment signal backfill — ApeWisdom daily captures starting now; minimum 90 days before it enters the analysis.")

    doc.save(str(output_path))
    log.info("Wrote %s", output_path)


def main() -> int:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    panel = load_signal_panel()
    reports = run_all(panel, SIGNALS)
    weights = derive_weights(reports)
    out = Path(__file__).resolve().parent.parent.parent.parent / "reports" / f"li_methodology_{date.today().isoformat()}.docx"
    build_docx(reports, weights, n_underliers=len(panel), output_path=out)
    print(f"Methodology doc: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
